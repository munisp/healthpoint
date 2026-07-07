"""
Document Generation Service — Full Production Implementation
Generates PDF and DOCX documents from templates for the HealthPoint IDR Platform.
Supports: GFE letters, IDR dispute filings, EOB documents, provider agreements,
          NSA notices, audit reports, and patient communications.
"""

import asyncio
import io
import logging

# ── Shared HealthPoint infrastructure ─────────────────────────────────────────
import sys, os as _os
_repo_root = _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
if _repo_root not in sys.path:
    sys.path.insert(0, _repo_root)
from backend.shared.database import fetch, fetchrow, execute, fetchval, transaction, bootstrap_schema, get_pool
from backend.shared.cache import get_client as get_redis_client, rate_limit_check, set_json, get_json
from backend.shared.auth import get_current_user, require_role, require_admin, require_provider, security_headers_middleware, TokenPayload
from backend.shared.messaging import publish, Topics
# ─────────────────────────────────────────────────────────────────────────────

import os
import uuid
from datetime import datetime, date
from enum import Enum
from typing import Any, Dict, List, Optional

import asyncpg
import boto3
from botocore.exceptions import ClientError
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import BackgroundTasks, FastAPI, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
from shared.telemetry import setup_telemetry, instrument_fastapi, get_tracer
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DATABASE_URL = os.environ["DATABASE_URL"]
S3_BUCKET = os.getenv("S3_BUCKET", "healthpoint-documents")
S3_REGION = os.getenv("AWS_REGION", "us-east-1")

setup_telemetry(service_name="document-generation-service", service_version="1.0.0")
app = FastAPI(
instrument_fastapi(app)

app.middleware("http")(security_headers_middleware)
    title="HealthPoint Document Generation Service",
    description="Generates PDF and DOCX documents from templates for IDR/NSA workflows.",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("ALLOWED_ORIGINS", "http://localhost:3000").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Enums & Models ──────────────────────────────────────────────────────────

class DocumentType(str, Enum):
    GFE_LETTER = "gfe_letter"
    IDR_FILING = "idr_filing"
    EOB = "explanation_of_benefits"
    PROVIDER_AGREEMENT = "provider_agreement"
    NSA_NOTICE = "nsa_notice"
    AUDIT_REPORT = "audit_report"
    PATIENT_COMMUNICATION = "patient_communication"
    DISPUTE_RESOLUTION_NOTICE = "dispute_resolution_notice"
    PAYMENT_DETERMINATION = "payment_determination"
    ARBITRATION_AWARD = "arbitration_award"


class OutputFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"


class PatientInfo(BaseModel):
    first_name: str
    last_name: str
    date_of_birth: Optional[date] = None
    member_id: Optional[str] = None
    address: Optional[str] = None


class ProviderInfo(BaseModel):
    name: str
    npi: Optional[str] = None
    tin: Optional[str] = None
    address: Optional[str] = None
    phone: Optional[str] = None
    specialty: Optional[str] = None


class ServiceItem(BaseModel):
    cpt_code: str
    description: str
    date_of_service: Optional[date] = None
    billed_amount: float
    allowed_amount: Optional[float] = None
    patient_responsibility: Optional[float] = None
    quantity: int = 1


class DocumentRequest(BaseModel):
    document_type: DocumentType
    output_format: OutputFormat = OutputFormat.PDF
    template_data: Dict[str, Any] = Field(default_factory=dict)
    patient: Optional[PatientInfo] = None
    provider: Optional[ProviderInfo] = None
    service_items: Optional[List[ServiceItem]] = None
    dispute_id: Optional[str] = None
    claim_id: Optional[str] = None
    reference_number: Optional[str] = None
    generated_by: Optional[str] = None
    store_to_s3: bool = True


class DocumentResponse(BaseModel):
    document_id: str
    document_type: str
    output_format: str
    s3_key: Optional[str] = None
    s3_url: Optional[str] = None
    file_size_bytes: int
    generated_at: datetime
    reference_number: Optional[str] = None


# ─── Database ────────────────────────────────────────────────────────────────

_pool: Optional[asyncpg.Pool] = None


async def get_db_pool() -> Optional[asyncpg.Pool]:
    global _pool
    if _pool is None:
        try:
            _pool = await asyncpg.create_pool(DATABASE_URL, min_size=2, max_size=10)
        except Exception as e:
            logger.warning(f"DB pool creation failed: {e}")
    return _pool


async def store_document_record(pool, document_id, doc_type, output_format,
                                 s3_key, file_size, claim_id, dispute_id, generated_by):
    if pool is None:
        return
    try:
        await pool.execute(
            """INSERT INTO documents (id, document_type, format, s3_key, file_size_bytes,
               claim_id, dispute_id, generated_by, created_at)
               VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9) ON CONFLICT (id) DO NOTHING""",
            document_id, doc_type, output_format, s3_key, file_size,
            claim_id, dispute_id, generated_by, datetime.utcnow(),
        )
    except Exception as e:
        logger.warning(f"Failed to store document record: {e}")


# ─── S3 Storage ──────────────────────────────────────────────────────────────

def upload_to_s3(content: bytes, key: str, content_type: str) -> Optional[str]:
    try:
        s3 = boto3.client("s3", region_name=S3_REGION)
        s3.put_object(Bucket=S3_BUCKET, Key=key, Body=content,
                      ContentType=content_type, ServerSideEncryption="AES256")
        return s3.generate_presigned_url("get_object",
                                          Params={"Bucket": S3_BUCKET, "Key": key},
                                          ExpiresIn=3600)
    except ClientError as e:
        logger.warning(f"S3 upload failed: {e}")
        return None


# ─── PDF Helpers ─────────────────────────────────────────────────────────────

def _header(story, styles, title, subtitle="", ref=""):
    h_style = ParagraphStyle("H", parent=styles["Normal"], fontSize=18,
                              textColor=colors.HexColor("#1a3a5c"), fontName="Helvetica-Bold")
    s_style = ParagraphStyle("S", parent=styles["Normal"], fontSize=11,
                              textColor=colors.HexColor("#4a6fa5"))
    story.append(Paragraph("HealthPoint IDR Platform", h_style))
    story.append(Paragraph(title, s_style))
    if subtitle:
        story.append(Paragraph(subtitle, s_style))
    if ref:
        story.append(Paragraph(f"Reference: {ref}", styles["Normal"]))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1a3a5c")))
    story.append(Spacer(1, 0.2 * inch))


def _info_table(data):
    tbl = Table(data, colWidths=[2.5 * inch, 4 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    return tbl


def _service_table(items):
    rows = [["CPT Code", "Description", "Date of Service", "Billed Amount", "Patient Resp."]]
    total_billed = total_resp = 0.0
    for item in items:
        resp = item.patient_responsibility or item.billed_amount
        rows.append([item.cpt_code, item.description[:40],
                     str(item.date_of_service) if item.date_of_service else "TBD",
                     f"${item.billed_amount:,.2f}", f"${resp:,.2f}"])
        total_billed += item.billed_amount
        total_resp += resp
    rows.append(["", "TOTAL", "", f"${total_billed:,.2f}", f"${total_resp:,.2f}"])
    tbl = Table(rows, colWidths=[1*inch, 2.5*inch, 1.2*inch, 1.3*inch, 1.2*inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f0f4f8")]),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#dce8f5")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    return tbl


# ─── PDF Generators ──────────────────────────────────────────────────────────

def generate_gfe_pdf(req: DocumentRequest) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    _header(story, styles, "GOOD FAITH ESTIMATE", "Notice of Expected Charges",
            req.reference_number or str(uuid.uuid4())[:8].upper())

    p = req.patient
    story.append(Paragraph("Patient Information", styles["Heading2"]))
    story.append(_info_table([
        ("Patient Name:", f"{p.first_name} {p.last_name}" if p else "N/A"),
        ("Date of Birth:", str(p.date_of_birth) if p and p.date_of_birth else "N/A"),
        ("Member ID:", p.member_id if p else "N/A"),
        ("GFE Date:", datetime.utcnow().strftime("%B %d, %Y")),
        ("Valid Through:", req.template_data.get("valid_through", "90 days from issue date")),
    ]))
    story.append(Spacer(1, 0.2 * inch))

    if req.provider:
        pv = req.provider
        story.append(Paragraph("Provider Information", styles["Heading2"]))
        story.append(_info_table([
            ("Provider Name:", pv.name),
            ("NPI:", pv.npi or "N/A"),
            ("Specialty:", pv.specialty or "N/A"),
            ("Address:", pv.address or "N/A"),
            ("Phone:", pv.phone or "N/A"),
        ]))
        story.append(Spacer(1, 0.2 * inch))

    if req.service_items:
        story.append(Paragraph("Expected Services and Costs", styles["Heading2"]))
        story.append(_service_table(req.service_items))
        story.append(Spacer(1, 0.2 * inch))

    notice = ParagraphStyle("N", parent=styles["Normal"], fontSize=8, textColor=colors.grey)
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "This Good Faith Estimate is provided pursuant to the No Surprises Act (Public Law 116-260). "
        "If you receive a bill that is at least $400 more than this GFE, you have the right to dispute "
        "the bill. Visit www.cms.gov/nosurprises for more information.", notice))
    doc.build(story)
    return buf.getvalue()


def generate_idr_filing_pdf(req: DocumentRequest) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    _header(story, styles, "INDEPENDENT DISPUTE RESOLUTION FILING",
            "No Surprises Act — Federal IDR Process",
            req.dispute_id or req.reference_number or "N/A")

    td = req.template_data
    story.append(Paragraph("Dispute Details", styles["Heading2"]))
    story.append(_info_table([
        ("Dispute ID:", req.dispute_id or "N/A"),
        ("Filing Date:", datetime.utcnow().strftime("%B %d, %Y")),
        ("Claim ID:", req.claim_id or "N/A"),
        ("Dispute Type:", td.get("dispute_type", "Out-of-Network Billing")),
        ("Initiating Party:", td.get("initiating_party", "N/A")),
        ("Responding Party:", td.get("responding_party", "N/A")),
        ("Disputed Amount:", f"${float(td.get('disputed_amount', 0)):,.2f}"),
        ("QPA Amount:", f"${float(td.get('qpa_amount', 0)):,.2f}"),
        ("IDR Entity:", td.get("idr_entity", "N/A")),
        ("Submission Deadline:", td.get("submission_deadline", "N/A")),
    ]))
    story.append(Spacer(1, 0.2 * inch))

    if req.patient:
        story.append(Paragraph("Patient Information", styles["Heading2"]))
        story.append(_info_table([
            ("Patient Name:", f"{req.patient.first_name} {req.patient.last_name}"),
            ("Member ID:", req.patient.member_id or "N/A"),
        ]))
        story.append(Spacer(1, 0.2 * inch))

    if td.get("supporting_rationale"):
        story.append(Paragraph("Supporting Rationale", styles["Heading2"]))
        story.append(Paragraph(td["supporting_rationale"], styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))

    if req.service_items:
        story.append(Paragraph("Disputed Services", styles["Heading2"]))
        story.append(_service_table(req.service_items))

    doc.build(story)
    return buf.getvalue()


def generate_eob_pdf(req: DocumentRequest) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    _header(story, styles, "EXPLANATION OF BENEFITS (EOB)",
            "Summary of Claim Processing", req.claim_id or "N/A")

    td = req.template_data
    p = req.patient
    story.append(Paragraph("Claim Summary", styles["Heading2"]))
    story.append(_info_table([
        ("Member Name:", f"{p.first_name} {p.last_name}" if p else "N/A"),
        ("Member ID:", p.member_id if p else "N/A"),
        ("Claim Number:", req.claim_id or "N/A"),
        ("Date Processed:", datetime.utcnow().strftime("%B %d, %Y")),
        ("Plan Name:", td.get("plan_name", "HealthPoint Insurance Plan")),
        ("Group Number:", td.get("group_number", "N/A")),
    ]))
    story.append(Spacer(1, 0.2 * inch))

    if req.service_items:
        story.append(Paragraph("Service Details", styles["Heading2"]))
        rows = [["CPT", "Description", "Billed", "Allowed", "Plan Paid", "Your Resp."]]
        total_b = total_a = total_p = total_r = 0.0
        for item in req.service_items:
            allowed = item.allowed_amount or item.billed_amount * 0.8
            resp = item.patient_responsibility or max(0, item.billed_amount - allowed)
            paid = allowed - resp
            rows.append([item.cpt_code, item.description[:28],
                         f"${item.billed_amount:,.2f}", f"${allowed:,.2f}",
                         f"${paid:,.2f}", f"${resp:,.2f}"])
            total_b += item.billed_amount; total_a += allowed
            total_p += paid; total_r += resp
        rows.append(["", "TOTALS", f"${total_b:,.2f}", f"${total_a:,.2f}",
                     f"${total_p:,.2f}", f"${total_r:,.2f}"])
        tbl = Table(rows, colWidths=[0.7*inch, 2.2*inch, 1*inch, 1*inch, 1*inch, 1.1*inch])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f0f4f8")]),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#dce8f5")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)
    doc.build(story)
    return buf.getvalue()


def generate_generic_pdf(req: DocumentRequest) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    title_map = {
        DocumentType.NSA_NOTICE: "NO SURPRISES ACT NOTICE",
        DocumentType.AUDIT_REPORT: "AUDIT COMPLIANCE REPORT",
        DocumentType.PATIENT_COMMUNICATION: "PATIENT COMMUNICATION",
        DocumentType.DISPUTE_RESOLUTION_NOTICE: "DISPUTE RESOLUTION NOTICE",
        DocumentType.PAYMENT_DETERMINATION: "PAYMENT DETERMINATION",
        DocumentType.ARBITRATION_AWARD: "ARBITRATION AWARD",
        DocumentType.PROVIDER_AGREEMENT: "PROVIDER AGREEMENT",
    }
    title = title_map.get(req.document_type, str(req.document_type).replace("_", " ").upper())
    _header(story, styles, title, "", req.reference_number or "N/A")

    td = req.template_data
    if td:
        story.append(Paragraph("Document Details", styles["Heading2"]))
        rows = [(str(k).replace("_", " ").title() + ":", str(v)) for k, v in td.items()]
        if rows:
            story.append(_info_table(rows))
            story.append(Spacer(1, 0.2 * inch))
    if td.get("body_text"):
        story.append(Paragraph(td["body_text"], styles["Normal"]))
    doc.build(story)
    return buf.getvalue()


# ─── DOCX Generator ──────────────────────────────────────────────────────────

def generate_docx(req: DocumentRequest) -> bytes:
    doc = DocxDocument()
    title_map = {
        DocumentType.GFE_LETTER: "GOOD FAITH ESTIMATE",
        DocumentType.IDR_FILING: "IDR DISPUTE FILING",
        DocumentType.EOB: "EXPLANATION OF BENEFITS",
        DocumentType.PROVIDER_AGREEMENT: "PROVIDER AGREEMENT",
        DocumentType.NSA_NOTICE: "NO SURPRISES ACT NOTICE",
        DocumentType.AUDIT_REPORT: "AUDIT COMPLIANCE REPORT",
        DocumentType.PATIENT_COMMUNICATION: "PATIENT COMMUNICATION",
        DocumentType.DISPUTE_RESOLUTION_NOTICE: "DISPUTE RESOLUTION NOTICE",
        DocumentType.PAYMENT_DETERMINATION: "PAYMENT DETERMINATION",
        DocumentType.ARBITRATION_AWARD: "ARBITRATION AWARD",
    }
    title = title_map.get(req.document_type, str(req.document_type).upper())
    h = doc.add_heading("HealthPoint IDR Platform", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading(title, 1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if req.reference_number:
        doc.add_paragraph(f"Reference: {req.reference_number}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}")

    if req.patient:
        doc.add_heading("Patient Information", 2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        for label, val in [
            ("Patient Name", f"{req.patient.first_name} {req.patient.last_name}"),
            ("Member ID", req.patient.member_id or "N/A"),
            ("Date of Birth", str(req.patient.date_of_birth) if req.patient.date_of_birth else "N/A"),
        ]:
            r = tbl.add_row().cells
            r[0].text = label; r[1].text = val

    if req.provider:
        doc.add_heading("Provider Information", 2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        for label, val in [
            ("Provider Name", req.provider.name),
            ("NPI", req.provider.npi or "N/A"),
            ("Specialty", req.provider.specialty or "N/A"),
        ]:
            r = tbl.add_row().cells
            r[0].text = label; r[1].text = val

    td = req.template_data
    if td:
        doc.add_heading("Details", 2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        for k, v in td.items():
            r = tbl.add_row().cells
            r[0].text = str(k).replace("_", " ").title()
            r[1].text = str(v)

    if req.service_items:
        doc.add_heading("Service Items", 2)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["CPT Code", "Description", "Date of Service", "Billed", "Patient Resp."]):
            hdr[i].text = h
        for item in req.service_items:
            r = tbl.add_row().cells
            r[0].text = item.cpt_code
            r[1].text = item.description
            r[2].text = str(item.date_of_service) if item.date_of_service else "TBD"
            r[3].text = f"${item.billed_amount:,.2f}"
            r[4].text = f"${item.patient_responsibility:,.2f}" if item.patient_responsibility else "N/A"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Core Logic ──────────────────────────────────────────────────────────────

def _generate_bytes(req: DocumentRequest) -> bytes:
    if req.output_format == OutputFormat.DOCX:
        return generate_docx(req)
    if req.document_type == DocumentType.GFE_LETTER:
        return generate_gfe_pdf(req)
    elif req.document_type == DocumentType.IDR_FILING:
        return generate_idr_filing_pdf(req)
    elif req.document_type == DocumentType.EOB:
        return generate_eob_pdf(req)
    else:
        return generate_generic_pdf(req)


async def _process(req: DocumentRequest) -> DocumentResponse:
    document_id = str(uuid.uuid4())
    content = _generate_bytes(req)
    file_size = len(content)
    ext = "pdf" if req.output_format == OutputFormat.PDF else "docx"
    content_type = "application/pdf" if req.output_format == OutputFormat.PDF else \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    s3_key = s3_url = None
    if req.store_to_s3:
        s3_key = f"documents/{req.document_type.value}/{document_id}.{ext}"
        s3_url = upload_to_s3(content, s3_key, content_type)

    pool = await get_db_pool()
    await store_document_record(pool, document_id, req.document_type.value,
                                 req.output_format.value, s3_key, file_size,
                                 req.claim_id, req.dispute_id, req.generated_by)
    return DocumentResponse(
        document_id=document_id,
        document_type=req.document_type.value,
        output_format=req.output_format.value,
        s3_key=s3_key,
        s3_url=s3_url,
        file_size_bytes=file_size,
        generated_at=datetime.utcnow(),
        reference_number=req.reference_number,
    )


# ─── API Endpoints ───────────────────────────────────────────────────────────

@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "document-generation", "version": "2.0.0"}


@app.post("/generate-pdf")
async def generate_pdf_legacy(html_content: str,
    current_user: TokenPayload = Depends(get_current_user),
):
    """Legacy endpoint: generate PDF from raw HTML using weasyprint."""
    try:
        from weasyprint import HTML
        pdf = HTML(string=html_content).write_pdf()
        return StreamingResponse(io.BytesIO(pdf), media_type="application/pdf")
    except ImportError:
        raise HTTPException(status_code=501, detail="weasyprint not available; use /api/v1/documents/generate")


@app.post("/api/v1/documents/generate", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def generate_document(req: DocumentRequest,
    current_user: TokenPayload = Depends(get_current_user),
):
    """Generate a document (PDF or DOCX) from a template."""
    try:
        return await _process(req)
    except Exception as e:
        logger.error(f"Document generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


@app.post("/api/v1/documents/generate/stream")
async def generate_document_stream(req: DocumentRequest,
    current_user: TokenPayload = Depends(get_current_user),
):
    """Generate and stream a document directly to the client."""
    try:
        req.store_to_s3 = False
        content = _generate_bytes(req)
        ext = "pdf" if req.output_format == OutputFormat.PDF else "docx"
        content_type = "application/pdf" if req.output_format == OutputFormat.PDF else \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{req.document_type.value}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.{ext}"
        return StreamingResponse(
            io.BytesIO(content), media_type=content_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"Document stream failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/documents/bulk", status_code=status.HTTP_202_ACCEPTED)
async def generate_bulk_documents(documents: List[DocumentRequest], background_tasks: BackgroundTasks,
    current_user: TokenPayload = Depends(get_current_user),
):
    """Generate multiple documents in bulk (async)."""
    if len(documents) > 100:
        raise HTTPException(status_code=400, detail="Maximum 100 documents per bulk request.")
    batch_id = str(uuid.uuid4())

    async def process_batch():
        for req in documents:
            try:
                await _process(req)
            except Exception as e:
                logger.error(f"Bulk doc failed: {e}")

    background_tasks.add_task(process_batch)
    return {"batch_id": batch_id, "document_count": len(documents), "status": "processing"}


@app.get("/api/v1/documents/{document_id}")
async def get_document_info(document_id: str,
    current_user: TokenPayload = Depends(get_current_user),
):
    """Retrieve metadata for a previously generated document."""
    pool = await get_db_pool()
    if pool is None:
        raise HTTPException(status_code=503, detail="Database unavailable")
    row = await pool.fetchrow("SELECT * FROM documents WHERE id = $1", document_id)
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return dict(row)


@app.get("/api/v1/documents")
async def list_documents(document_type: Optional[str] = None, claim_id: Optional[str] = None,
                          dispute_id: Optional[str] = None, limit: int = 50, offset: int = 0,
                              current_user: TokenPayload = Depends(get_current_user),
                          ):
    """List documents with optional filtering."""
    pool = await get_db_pool()
    if pool is None:
        raise HTTPException(status_code=503, detail="Database unavailable")
    conditions, params, idx = [], [], 1
    for col, val in [("document_type", document_type), ("claim_id", claim_id), ("dispute_id", dispute_id)]:
        if val:
            conditions.append(f"{col} = ${idx}")
            params.append(val)
            idx += 1
    where = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    params.extend([limit, offset])
    rows = await pool.fetch(
        f"SELECT * FROM documents {where} ORDER BY created_at DESC LIMIT ${idx} OFFSET ${idx+1}", *params)
    return {"documents": [dict(r) for r in rows], "total": len(rows)}


@app.get("/api/v1/documents/types/supported")
async def list_supported_types(,
    current_user: TokenPayload = Depends(get_current_user),
):
    return {"document_types": [t.value for t in DocumentType],
            "output_formats": [f.value for f in OutputFormat]}


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8030"))
    uvicorn.run(app, host="0.0.0.0", port=port)